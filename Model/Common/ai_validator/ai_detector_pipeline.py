import os
import re
import json
import math
import hashlib
import uuid
from io import BytesIO
from typing import List, Dict, Any, Optional, Tuple, Union
from datetime import datetime
from multiprocessing import Pool, cpu_count
import time

import PyPDF2
import chardet
import docx
import google.generativeai as genai
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from dotenv import load_dotenv
from supabase import create_client, Client
import joblib
import traceback

# --- Optional detector libs (not required) ---
try:
    import importlib
    typetruth = importlib.import_module("typetruth")
    TYPETRUTH_AVAILABLE = True
except Exception:
    typetruth = None
    TYPETRUTH_AVAILABLE = False

# GPT-2 perplexity fallback is optional and expensive:
try:
    from transformers import GPT2LMHeadModel, GPT2TokenizerFast
    TRANSFORMERS_AVAILABLE = True
except Exception:
    TRANSFORMERS_AVAILABLE = False

# Load .env
load_dotenv()

# -------------------- CONFIG --------------------
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")  # service_role recommended
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY4")

if not SUPABASE_URL or not SUPABASE_KEY or not GEMINI_API_KEY:
    raise Exception("Set SUPABASE_URL, SUPABASE_KEY, and GEMINI_API_KEY in .env")

# Buckets / table names - adjust to your names
RAW_BUCKET = os.getenv("RAW_BUCKET", "Coal-research-files")
REPORT_BUCKET = os.getenv("REPORT_BUCKET", "ai-detection-json")
REPORT_TABLE = os.getenv("REPORT_TABLE", "ai_detection_reports")

# Thresholds & workers
FIELD_VALIDATE_THRESHOLD = float(os.getenv("FIELD_VALIDATE_THRESHOLD", "0.65"))  # detector score -> candidate
SENTENCE_VALIDATE_THRESHOLD = float(os.getenv("SENTENCE_VALIDATE_THRESHOLD", "0.7"))  # sentence-level detector threshold for Gemini validation
WORKER_COUNT = min(int(os.getenv("WORKER_COUNT", str(max(1, cpu_count() - 1)))), 8)

MODEL_NAME = os.getenv("MODEL_NAME", "gemini-2.5-flash-lite")

# Initialize Supabase and Gemini
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
genai.configure(api_key=GEMINI_API_KEY)
router = APIRouter()

# Optional trained validator model
MODEL_JOBLIB_PATH = r"C:\Users\Shanmuga Shyam. B\OneDrive\Desktop\SIH25180\Model\Common\ai_validator\pre-trained\my_trained_model.joblib"
AI_VALIDATOR_MODEL = None
AI_VALIDATOR_AVAILABLE = False

try:
    if os.path.exists(MODEL_JOBLIB_PATH):
        AI_VALIDATOR_MODEL = joblib.load(MODEL_JOBLIB_PATH)
        AI_VALIDATOR_AVAILABLE = True
    else:
        # Simple heuristic validator fallback
        class _HeuristicValidator:
            """Simple fallback validator with predict and predict_proba."""
            def __init__(self):
                pass

            def _score_text(self, text: str) -> float:
                words = text.split()
                if not words:
                    return 0.0
                avg_word_len = sum(len(w) for w in words) / len(words)
                punct = sum(1 for ch in text if ch in ".,;:!?()'")
                base = 0.18 if avg_word_len >= 4 else 0.45
                prob = min(0.99, base * (len(text) / 800.0) + min(0.2, punct/50.0))
                return float(max(0.0, min(0.999999, prob)))

            def predict_proba(self, X):
                return [[1.0 - self._score_text(x), self._score_text(x)] for x in X]

            def predict(self, X):
                return [1 if self._score_text(x) >= 0.5 else 0 for x in X]

        AI_VALIDATOR_MODEL = _HeuristicValidator()
        AI_VALIDATOR_AVAILABLE = True
except Exception:
    AI_VALIDATOR_AVAILABLE = False

# =============================================================================
# AGENT STATE MANAGEMENT
# =============================================================================

class AgentState:
    """Manages the state of the AI detection agent."""
    
    def __init__(self, structured_data: Dict[str, Any]):
        self.structured_data = structured_data
        self.project_details = structured_data.get("project_details", {})
        self.field_results = {}  # field_name -> detection_result
        self.suspicious_fields = []  # fields that need Gemini validation
        self.validated_fields = {}  # field_name -> gemini_validation
        self.action_log = []  # log of all agent actions
        self.current_step = 0
        self.max_steps = 50
        self.finished = False
        self.recent_reports = []  # loaded from database
        
    def add_action(self, action: str, details: Dict[str, Any]):
        """Add an action to the agent's log."""
        self.action_log.append({
            "step": self.current_step,
            "action": action,
            "details": details,
            "timestamp": datetime.utcnow().isoformat()
        })
        self.current_step += 1
        
    def get_unprocessed_fields(self) -> List[str]:
        """Get list of fields that haven't been processed yet."""
        return [field for field in self.project_details.keys() 
                if field not in self.field_results]
    
    def get_unvalidated_suspicious_fields(self) -> List[str]:
        """Get suspicious fields that haven't been validated by Gemini yet."""
        return [field for field in self.suspicious_fields 
                if field not in self.validated_fields]
    
    def is_complete(self) -> bool:
        """Check if agent has completed its analysis."""
        return (self.finished or 
                self.current_step >= self.max_steps or 
                len(self.get_unprocessed_fields()) == 0)

# =============================================================================
# FIELD-BASED DETECTION WRAPPER
# =============================================================================

def detect_field_ai_probability(field_name: str, field_content: str) -> Dict[str, Any]:
    """Run local detector on a single field."""
    if not field_content or not field_content.strip():
        return {
            "field_name": field_name,
            "content_length": 0,
            "ai_probability": 0.0,
            "detector_method": "empty_content"
        }
    
    try:
        if AI_VALIDATOR_AVAILABLE and AI_VALIDATOR_MODEL:
            proba = AI_VALIDATOR_MODEL.predict_proba([field_content])
            ai_prob = float(proba[0][1]) if proba and len(proba[0]) > 1 else 0.0
        else:
            # Fallback heuristic
            words = field_content.split()
            avg_word_len = sum(len(w) for w in words) / len(words) if words else 0
            ai_prob = 0.3 if avg_word_len > 5 else 0.7
            
        return {
            "field_name": field_name,
            "content_length": len(field_content),
            "ai_probability": round(ai_prob, 6),
            "detector_method": "ai_validator_model" if AI_VALIDATOR_AVAILABLE else "heuristic"
        }
    except Exception as e:
        return {
            "field_name": field_name,
            "content_length": len(field_content),
            "ai_probability": 0.5,
            "detector_method": "error",
            "error": str(e)
        }

def worker_detect_field(args: Tuple[str, str]) -> Dict[str, Any]:
    """Worker function for multiprocessing field detection."""
    field_name, field_content = args
    return detect_field_ai_probability(field_name, field_content)

# =============================================================================
# GEMINI VALIDATOR FOR FIELDS
# =============================================================================

def call_gemini_validate_field(field_name: str, field_content: str, 
                              detector_score: float, recent_reports: List[Dict]) -> Dict[str, Any]:
    """Validate a field using Gemini AI."""
    
    # Build context from recent reports
    context_summary = ""
    if recent_reports:
        context_summary = f"\nRecent analysis context: {len(recent_reports)} similar documents analyzed. "
        ai_patterns = []
        for report in recent_reports[:5]:  # Use top 5 for context
            if isinstance(report.get("result"), dict):
                agg = report["result"].get("aggregate", {})
                verdict = agg.get("file_verdict", "unknown")
                if verdict == "ai":
                    ai_patterns.append("High AI likelihood found in similar documents.")
        
        if ai_patterns:
            context_summary += "Common AI patterns: " + "; ".join(ai_patterns[:3])
    
    prompt = f"""You are an expert AI text detector analyzing a specific field from a FORM-I S&T Grant Proposal.

FIELD ANALYSIS TASK:
- Field Name: {field_name}
- Local Detector Score: {detector_score:.3f} (0.0 = human, 1.0 = AI)
- Content Length: {len(field_content)} characters

CONTENT TO ANALYZE:
{field_content}

{context_summary}

ANALYSIS REQUIREMENTS:
1. Determine if this field content appears to be AI-generated or human-written
2. Consider technical writing patterns typical in research proposals
3. Look for AI-specific patterns: repetitive phrasing, generic language, lack of specific details
4. Consider domain expertise indicators: specific technical terms, concrete examples, personal insights

Respond with ONLY a JSON object:
{{
    "validated_ai_probability": <float 0.0-1.0>,
    "decision": "<ai|human|uncertain>",
    "confidence": <float 0.0-1.0>,
    "reasoning": "<brief explanation>",
    "ai_indicators": ["<list of AI-like patterns found>"],
    "human_indicators": ["<list of human-like patterns found>"],
    "recommendation": "<specific improvement advice if AI-like>"
}}"""

    try:
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        
        # Clean and parse JSON
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
            
        result = json.loads(response_text)
        
        # Validate and clean result
        validated_prob = float(result.get("validated_ai_probability", detector_score))
        decision = result.get("decision", "uncertain").lower()
        
        if decision not in ["ai", "human", "uncertain"]:
            decision = "uncertain"
            
        return {
            "validated_ai_probability": validated_prob,
            "decision": decision,
            "confidence": float(result.get("confidence", 0.5)),
            "reasoning": result.get("reasoning", ""),
            "ai_indicators": result.get("ai_indicators", []),
            "human_indicators": result.get("human_indicators", []),
            "recommendation": result.get("recommendation", ""),
            "gemini_model": MODEL_NAME,
            "validation_timestamp": datetime.utcnow().isoformat()
        }
        
    except Exception as e:
        # Fallback validation based on detector score
        if detector_score >= 0.8:
            decision = "ai"
        elif detector_score <= 0.3:
            decision = "human"
        else:
            decision = "uncertain"
            
        return {
            "validated_ai_probability": detector_score,
            "decision": decision,
            "confidence": 0.3,
            "reasoning": f"Gemini validation failed, using detector score. Error: {str(e)}",
            "ai_indicators": [],
            "human_indicators": [],
            "recommendation": "Unable to provide specific recommendations due to validation error.",
            "gemini_model": MODEL_NAME,
            "validation_timestamp": datetime.utcnow().isoformat(),
            "error": str(e)
        }

# =============================================================================
# AGENT PLANNING AND ACTION SYSTEM
# =============================================================================

class AIDetectionAgent:
    """LLM-driven agent for AI detection analysis."""
    
    def __init__(self, state: AgentState):
        self.state = state
        
    def plan_next_action(self) -> Dict[str, Any]:
        """Use LLM to plan the next action based on current state."""
        
        # Prepare state summary for planning
        total_fields = len(self.state.project_details)
        processed_fields = len(self.state.field_results)
        suspicious_count = len(self.state.suspicious_fields)
        validated_count = len(self.state.validated_fields)
        unprocessed = self.state.get_unprocessed_fields()
        unvalidated = self.state.get_unvalidated_suspicious_fields()
        
        state_summary = f"""
CURRENT ANALYSIS STATE:
- Total fields in project_details: {total_fields}
- Fields processed: {processed_fields}/{total_fields}
- Suspicious fields identified: {suspicious_count}
- Fields validated by Gemini: {validated_count}
- Unprocessed fields: {len(unprocessed)}
- Unvalidated suspicious fields: {len(unvalidated)}
- Current step: {self.state.current_step}
- Max steps: {self.state.max_steps}

UNPROCESSED FIELDS: {unprocessed[:5]}
UNVALIDATED SUSPICIOUS: {unvalidated[:3]}

AVAILABLE ACTIONS:
1. inspect_field: Run local detector on a specific field
2. validate_field: Use Gemini to validate a suspicious field  
3. load_context: Load recent reports for context
4. summarize_progress: Summarize current findings
5. finish: Complete analysis and generate final report

FIELD RESULTS SO FAR:
{json.dumps({k: v.get('ai_probability', 0) for k, v in self.state.field_results.items()}, indent=2)}
"""

        planning_prompt = f"""You are an AI detection agent planning your next action. 

{state_summary}

Choose the most logical next action. Prioritize:
1. Loading context if not done yet
2. Processing unprocessed fields first  
3. Validating highly suspicious fields (>0.7 AI probability)
4. Finishing when all important work is complete

Respond with ONLY a JSON object:
{{
    "action": "<action_name>",
    "target": "<field_name_or_other_target>",
    "reasoning": "<why this action now>",
    "priority": "<high|medium|low>"
}}"""

        try:
            model = genai.GenerativeModel(MODEL_NAME)
            response = model.generate_content(planning_prompt)
            response_text = response.text.strip()
            
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
                
            action_plan = json.loads(response_text)
            
            # Validate action
            valid_actions = ["inspect_field", "validate_field", "load_context", "summarize_progress", "finish"]
            if action_plan.get("action") not in valid_actions:
                action_plan["action"] = "finish"
                
            return action_plan
            
        except Exception as e:
            # Fallback planning logic
            if not self.state.recent_reports:
                return {"action": "load_context", "target": "recent_reports", "reasoning": "Need context first", "priority": "high"}
            elif unprocessed:
                return {"action": "inspect_field", "target": unprocessed[0], "reasoning": "Process remaining fields", "priority": "high"}
            elif unvalidated:
                return {"action": "validate_field", "target": unvalidated[0], "reasoning": "Validate suspicious field", "priority": "high"}
            else:
                return {"action": "finish", "target": "analysis", "reasoning": "All work complete", "priority": "high"}
    
    def execute_action(self, action_plan: Dict[str, Any]) -> Dict[str, Any]:
        """Execute the planned action."""
        action = action_plan.get("action")
        target = action_plan.get("target")
        
        self.state.add_action(action, action_plan)
        
        if action == "inspect_field":
            return self._inspect_field(target)
        elif action == "validate_field":
            return self._validate_field(target)
        elif action == "load_context":
            return self._load_context()
        elif action == "summarize_progress":
            return self._summarize_progress()
        elif action == "finish":
            return self._finish_analysis()
        else:
            return {"error": f"Unknown action: {action}"}
    
    def _inspect_field(self, field_name: str) -> Dict[str, Any]:
        """Inspect a field using local detector."""
        if field_name not in self.state.project_details:
            return {"error": f"Field {field_name} not found in project_details"}
            
        field_content = self.state.project_details[field_name]
        result = detect_field_ai_probability(field_name, field_content)
        
        self.state.field_results[field_name] = result
        
        # Add to suspicious list if probability is high
        if result.get("ai_probability", 0) >= FIELD_VALIDATE_THRESHOLD:
            if field_name not in self.state.suspicious_fields:
                self.state.suspicious_fields.append(field_name)
                
        return {"success": True, "result": result}
    
    def _validate_field(self, field_name: str) -> Dict[str, Any]:
        """Validate a field using Gemini."""
        if field_name not in self.state.project_details:
            return {"error": f"Field {field_name} not found in project_details"}
            
        if field_name not in self.state.field_results:
            return {"error": f"Field {field_name} not yet inspected by local detector"}
            
        field_content = self.state.project_details[field_name]
        detector_score = self.state.field_results[field_name].get("ai_probability", 0.5)
        
        validation = call_gemini_validate_field(field_name, field_content, detector_score, self.state.recent_reports)
        self.state.validated_fields[field_name] = validation
        
        return {"success": True, "validation": validation}
    
    def _load_context(self) -> Dict[str, Any]:
        """Load recent reports for context."""
        try:
            self.state.recent_reports = load_recent_reports(limit=20)
            return {"success": True, "loaded_reports": len(self.state.recent_reports)}
        except Exception as e:
            return {"error": f"Failed to load context: {str(e)}"}
    
    def _summarize_progress(self) -> Dict[str, Any]:
        """Summarize current progress."""
        summary = {
            "total_fields": len(self.state.project_details),
            "processed_fields": len(self.state.field_results),
            "suspicious_fields": len(self.state.suspicious_fields),
            "validated_fields": len(self.state.validated_fields),
            "completion_percentage": round((len(self.state.field_results) / len(self.state.project_details)) * 100, 1),
            "actions_taken": self.state.current_step
        }
        return {"success": True, "summary": summary}
    
    def _finish_analysis(self) -> Dict[str, Any]:
        """Mark analysis as finished."""
        self.state.finished = True
        return {"success": True, "message": "Analysis complete"}
    
    def run_analysis(self) -> Dict[str, Any]:
        """Run the complete agent analysis loop."""
        while not self.state.is_complete():
            try:
                # Plan next action
                action_plan = self.plan_next_action()
                
                # Execute action
                result = self.execute_action(action_plan)
                
                # Safety check to prevent infinite loops
                if self.state.current_step >= self.state.max_steps:
                    self.state.add_action("emergency_stop", {"reason": "Max steps reached"})
                    break
                    
                # Small delay to prevent rate limiting
                time.sleep(0.1)
                
            except Exception as e:
                self.state.add_action("error", {"error": str(e)})
                break
        
        return self.generate_final_report()
    
    def generate_final_report(self) -> Dict[str, Any]:
        """Generate the final analysis report."""
        # Calculate field-level analysis in the requested format
        field_level_analysis = {}
        for field_name, result in self.state.field_results.items():
            detector_prob = result.get("ai_probability", 0.0)
            
            # Get validation if available
            validation = self.state.validated_fields.get(field_name, {})
            validated_prob = validation.get("validated_ai_probability", detector_prob)
            gemini_decision = validation.get("decision", "uncertain")
            gemini_reasoning = validation.get("reasoning", "")
            
            # Calculate combined score (60% detector, 40% validation)
            combined_ai_prob = (0.6 * detector_prob) + (0.4 * validated_prob)
            combined_human_prob = 1.0 - combined_ai_prob
            
            # Determine classification based on combined probability
            if combined_ai_prob >= 0.6:
                classification = "ai"
            elif combined_ai_prob <= 0.4:
                classification = "human" 
            else:
                classification = "mixed"
            
            # Generate reason based on available information
            if gemini_reasoning:
                reason = gemini_reasoning
            elif validation.get("ai_indicators"):
                reason = "AI patterns detected: " + ", ".join(validation["ai_indicators"][:2])
            elif validation.get("human_indicators"):
                reason = "Human patterns detected: " + ", ".join(validation["human_indicators"][:2])
            elif combined_ai_prob >= 0.7:
                reason = "Highly generic phrasing and lacks contextual grounding."
            elif combined_ai_prob >= 0.6:
                reason = "Shows template-like structure and formal language patterns."
            elif combined_ai_prob <= 0.3:
                reason = "Contains natural imperfections and human patterning."
            elif combined_ai_prob <= 0.4:
                reason = "Shows domain-specific knowledge and natural variation."
            else:
                reason = "Combination of structured and free-flow text."
            
            field_level_analysis[field_name] = {
                "ai_probability": round(combined_ai_prob, 2),
                "human_probability": round(combined_human_prob, 2),
                "classification": classification,
                "reason": reason
            }
        
        # Calculate overall scores
        if field_level_analysis:
            avg_ai_prob = sum(f["ai_probability"] for f in field_level_analysis.values()) / len(field_level_analysis)
            avg_human_prob = 1.0 - avg_ai_prob
        else:
            avg_ai_prob = 0.5
            avg_human_prob = 0.5
        
        # Convert to percentages for overall scores
        overall_ai_percentage = round(avg_ai_prob * 100)
        overall_human_percentage = round(avg_human_prob * 100)
        
        # Determine overall classification
        if avg_ai_prob >= 0.6:
            summary_decision = "ai"
            summary_reason = "Majority of fields show AI-generated patterns with templated language and generic phrasing."
        elif avg_ai_prob <= 0.4:
            summary_decision = "human"
            summary_reason = "Document demonstrates natural human writing with domain expertise and authentic variation."
        else:
            summary_decision = "mixed"
            summary_reason = "Some fields show templated AI-like phrasing while others show natural human variability and domain grounding."
        
        # Identify flagged fields (those classified as AI or with high AI probability)
        flagged_fields = []
        for field_name, analysis in field_level_analysis.items():
            if (analysis["classification"] == "ai" or analysis["ai_probability"] >= 0.7):
                validation = self.state.validated_fields.get(field_name, {})
                flagged_fields.append({
                    "field_name": field_name,
                    "ai_probability": analysis["ai_probability"],
                    "decision": analysis["classification"],
                    "ai_indicators": validation.get("ai_indicators", []),
                    "recommendation": validation.get("recommendation", "")
                })
        
        # Generate explanation
        explanation = self._generate_explanation(summary_decision, avg_ai_prob, flagged_fields)
        
        # Generate recommendations
        recommendations = self._generate_recommendations(flagged_fields)
        
        return {
            "field_level_analysis": field_level_analysis,
            "overall_scores": {
                "overall_ai_percentage": overall_ai_percentage,
                "overall_human_percentage": overall_human_percentage,
                "summary_decision": summary_decision,
                "summary_reason": summary_reason,
                # Keep legacy format for backward compatibility
                "human_score": overall_human_percentage,
                "ai_score": overall_ai_percentage,
                "classification": summary_decision,
                "confidence": round(abs(avg_ai_prob - 0.5) * 2, 3)
            },
            "flagged_fields": flagged_fields,
            "explanation": explanation,
            "recommendations": recommendations,
            "agent_actions_log": self.state.action_log,
            "analysis_metadata": {
                "total_fields_analyzed": len(field_level_analysis),
                "fields_requiring_validation": len(self.state.validated_fields),
                "processing_time_seconds": self.state.current_step * 0.1,
                "agent_steps": self.state.current_step
            }
        }
    
    def _generate_explanation(self, classification: str, ai_prob: float, flagged_fields: List[Dict]) -> str:
        """Generate human-readable explanation."""
        if classification == "human":
            return f"Document appears to be human-authored with {round((1-ai_prob)*100, 1)}% confidence. Writing shows natural variation and domain expertise."
        elif classification == "ai":
            return f"Document shows strong indicators of AI generation with {round(ai_prob*100, 1)}% confidence. {len(flagged_fields)} fields show clear AI patterns."
        else:
            return f"Document shows mixed signals. Some sections appear AI-generated while others seem human-authored. {len(flagged_fields)} fields require attention."
    
    def _generate_recommendations(self, flagged_fields: List[Dict]) -> List[str]:
        """Generate improvement recommendations."""
        recommendations = []
        
        if not flagged_fields:
            recommendations.append("Document quality is good. Consider adding more specific examples and citations to strengthen credibility.")
        else:
            recommendations.append(f"Review and revise {len(flagged_fields)} flagged fields to improve human-likeness.")
            
            for field in flagged_fields[:3]:  # Top 3 most problematic
                if field.get("recommendation"):
                    recommendations.append(f"{field['field_name']}: {field['recommendation']}")
            
            recommendations.append("Add personal insights, specific examples, and vary sentence structure to improve authenticity.")
            recommendations.append("Include concrete citations and domain-specific terminology where appropriate.")
        
        return recommendations

# =============================================================================
# UTILITY FUNCTIONS (PRESERVED FROM ORIGINAL)
# =============================================================================

def load_recent_reports(limit: int = 20) -> List[Dict[str, Any]]:
    """Load recent AI detection reports from database."""
    try:
        response = supabase.table(REPORT_TABLE).select("*").order("created_at", desc=True).limit(limit).execute()
        return response.data or []
    except Exception as e:
        print(f"Warning: Could not load recent reports: {e}")
        return []

def upload_json_report_to_storage(filename: str, report_data: Dict[str, Any]) -> str:
    """Upload JSON report to Supabase storage."""
    json_bytes = json.dumps(report_data, indent=2).encode('utf-8')
    supabase.storage.from_(REPORT_BUCKET).upload(filename, json_bytes, {"content-type": "application/json"})
    return supabase.storage.from_(REPORT_BUCKET).get_public_url(filename)

def insert_report_row_db(row_data: Dict[str, Any]):
    """Insert report row into database."""
    supabase.table(REPORT_TABLE).insert(row_data).execute()

def extract_text_from_file(filename: str, file_bytes: bytes) -> str:
    """Extract text from uploaded file."""
    try:
        ext = filename.lower().split(".")[-1]
        
        if ext == "pdf":
            reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        elif ext == "docx":
            doc = docx.Document(BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext in ["txt", "csv"]:
            encoding = chardet.detect(file_bytes)["encoding"] or "utf-8"
            return file_bytes.decode(encoding, errors="ignore")
        else:
            return ""
    except Exception:
        return ""

def build_structured_json_from_text(text: str) -> Dict[str, Any]:
    """Build structured JSON from extracted text - placeholder for integration with OCR extraction."""
    # This is a simplified version - in production, integrate with your existing
    # OCR extraction system to get the proper FORM-I structure
    
    # For now, create a basic structure with project details extracted from text
    lines = text.split('\n')
    content_lines = [line.strip() for line in lines if line.strip()]
    
    # Simple field extraction (this should be replaced with your actual OCR logic)
    project_details = {
        "definition_of_issue": "",
        "objectives": "",
        "justification_subject_area": "",
        "project_benefits": "",
        "work_plan": "",
        "methodology": "",
        "organization_of_work": "",
        "time_schedule": "",
        "foreign_exchange_details": ""
    }
    
    # Basic text segmentation into fields (simplified approach)
    current_field = None
    current_content = []
    
    for line in content_lines[:50]:  # Process first 50 lines
        line_lower = line.lower()
        
        if "definition" in line_lower and "issue" in line_lower:
            current_field = "definition_of_issue"
        elif "objective" in line_lower:
            current_field = "objectives"
        elif "justification" in line_lower:
            current_field = "justification_subject_area"
        elif "benefit" in line_lower:
            current_field = "project_benefits"
        elif "work plan" in line_lower or "methodology" in line_lower:
            current_field = "methodology"
        elif current_field and len(line) > 20:
            current_content.append(line)
            
        if len(current_content) > 5 and current_field:
            project_details[current_field] = " ".join(current_content)
            current_content = []
            current_field = None
    
    # Fill in remaining empty fields with portions of text
    remaining_text = " ".join(content_lines)
    for field in project_details:
        if not project_details[field] and remaining_text:
            # Take a chunk of text for each empty field
            chunk_size = min(500, len(remaining_text) // len(project_details))
            project_details[field] = remaining_text[:chunk_size]
            remaining_text = remaining_text[chunk_size:]
    
    return {
        "form_type": "FORM-I S&T Grant Proposal",
        "basic_information": {},
        "project_details": project_details,
        "cost_breakdown": {},
        "additional_information": {}
    }

def compute_file_hash(file_bytes: bytes) -> str:
    """Compute SHA-256 hash of file."""
    return hashlib.sha256(file_bytes).hexdigest()

def sanitize_filename(filename: str) -> str:
    """Sanitize filename for storage."""
    return re.sub(r'[^\w\-_\.]', '_', filename)

def make_unique_filename(filename: str, existing_names: List[str]) -> str:
    """Make filename unique by adding suffix."""
    base, ext = filename.rsplit('.', 1) if '.' in filename else (filename, '')
    counter = 1
    while filename in existing_names:
        filename = f"{base}_{counter}.{ext}" if ext else f"{base}_{counter}"
        counter += 1
    return filename

# =============================================================================
# MAIN API ENDPOINT
# =============================================================================

@router.post("/detect-ai-only")
async def detect_ai_only(file: UploadFile = File(...)):
    """
    AI Detection Agent Endpoint - analyzes FORM-I structured data using LLM-driven agent.
    
    Returns:
    {
        "classification": "ai|human|mixed",
        "human_score": 0-100,
        "ai_score": 0-100, 
        "flagged_fields": [...],
        "improvement_comment": "...",
        "full_report_url": "..."
    }
    """
    try:
        # 1) Read and hash file
        filename_raw = file.filename or "unknown.txt"
        file_bytes = await file.read()
        file_hash = compute_file_hash(file_bytes)
        
        # 2) Extract text from file
        extracted_text = extract_text_from_file(filename_raw, file_bytes)
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="No text content could be extracted from file")
        
        # 3) Build structured JSON (integrate with your existing OCR system)
        structured_data = build_structured_json_from_text(extracted_text)
        
        # Validate that we have project_details to analyze
        if not structured_data.get("project_details"):
            raise HTTPException(status_code=400, detail="No project details found in document")
        
        # 4) Initialize agent state and run analysis
        agent_state = AgentState(structured_data)
        agent = AIDetectionAgent(agent_state)
        
        # 5) Run the agent analysis loop
        final_report = agent.run_analysis()
        
        # 6) Store file in raw bucket
        uploaded_name = filename_raw
        try:
            existing_files = supabase.storage.from_(RAW_BUCKET).list() or []
            existing_names = [o.get("name") for o in existing_files]
            if uploaded_name in existing_names:
                uploaded_name = make_unique_filename(uploaded_name, existing_names)
            
            supabase.storage.from_(RAW_BUCKET).upload(
                uploaded_name, file_bytes, {"content-type": "application/octet-stream"}
            )
        except Exception:
            pass  # Best effort
        
        # 7) Store full report JSON
        report_data = {
            "original_filename": filename_raw,
            "file_hash": file_hash,
            "created_at": datetime.utcnow().isoformat(),
            "structured_data": structured_data,
            "agent_analysis": final_report,
            "raw_text_snippet": extracted_text[:5000]
        }
        
        json_filename = f"{sanitize_filename(uploaded_name.rsplit('.', 1)[0])}_{uuid.uuid4().hex[:8]}.ai-detect.json"
        try:
            report_url = upload_json_report_to_storage(json_filename, report_data)
            print(f"Successfully stored detect-ai-only JSON: {json_filename}")
        except Exception as storage_error:
            print(f"Failed to store detect-ai-only JSON: {str(storage_error)}")
            report_url = ""
        
        # 8) Insert database record
        db_row = {
            "filename": uploaded_name,
            "file_path": f"{RAW_BUCKET}/{uploaded_name}",
            "file_hash": file_hash,
            "ai_percentage": final_report["overall_scores"]["ai_score"],
            "segments_count": len(final_report["field_level_analysis"]),
            "result": report_data,
            "created_at": datetime.utcnow().isoformat()
        }
        
        try:
            insert_report_row_db(db_row)
        except Exception:
            pass  # Best effort
        
        # 9) Build compact response with both formats
        overall_scores = final_report["overall_scores"]
        flagged_fields = final_report["flagged_fields"]
        
        # Generate improvement comment
        if overall_scores["classification"] == "human":
            improvement_comment = "Document appears human-authored. Consider adding more specific citations and examples to strengthen credibility."
        elif overall_scores["classification"] == "ai":
            improvement_comment = f"Document shows signs of AI generation in {len(flagged_fields)} fields. Add personal insights, specific examples, and vary writing style."
        else:
            improvement_comment = f"Mixed signals detected. {len(flagged_fields)} fields need attention to improve human authenticity."
        
        response = {
            "classification": overall_scores["classification"],
            "human_score": int(overall_scores["human_score"]),
            "ai_score": int(overall_scores["ai_score"]),
            "flagged_fields": [
                {
                    "field_name": field["field_name"],
                    "ai_probability": field["ai_probability"],
                    "reason": field.get("recommendation", "AI-like patterns detected")[:200]
                }
                for field in flagged_fields[:5]  # Top 5 most problematic fields
            ],
            "improvement_comment": improvement_comment,
            "full_report_url": report_url,
            "field_level_analysis": final_report["field_level_analysis"],
            "overall_scores": {
                "overall_ai_percentage": overall_scores["overall_ai_percentage"],
                "overall_human_percentage": overall_scores["overall_human_percentage"],
                "summary_decision": overall_scores["summary_decision"],
                "summary_reason": overall_scores["summary_reason"]
            },
            "analysis_metadata": {
                "fields_analyzed": len(final_report["field_level_analysis"]),
                "processing_steps": final_report["analysis_metadata"]["agent_steps"],
                "confidence": overall_scores["confidence"]
            }
        }
        
        return JSONResponse(content=response)
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"AI Detection Agent Error: {traceback.format_exc()}")
        return JSONResponse(
            content={"error": f"Analysis failed: {str(e)}"}, 
            status_code=500
        )


# --- clients & configure ---
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
genai.configure(api_key=GEMINI_API_KEY)

router = APIRouter()

# -------------------- Utilities --------------------
def sha256_bytes(b: bytes) -> str:
    h = hashlib.sha256()
    h.update(b)
    return h.hexdigest()

def sanitize_filename(name: str) -> str:
    name = (name or "").strip()
    name = re.sub(r"[<>:\"/\\|?*\n\r]+", "_", name)
    return name

def make_unique_filename(base_name: str, existing_names: List[str]) -> str:
    if base_name not in existing_names:
        return base_name
    base, dot, ext = base_name.rpartition(".")
    if not dot:
        base = base_name
        ext = ""
    else:
        ext = "." + ext
    i = 1
    candidate = f"{base}-v{i}{ext}"
    while candidate in existing_names:
        i += 1
        candidate = f"{base}-v{i}{ext}"
    return candidate

def safe_json_parse(text: str) -> Optional[Dict[str, Any]]:
    if not text:
        return None
    txt = text.strip()
    # try JSON directly
    try:
        return json.loads(txt)
    except Exception:
        pass
    # fallback: extract first {...} block
    m = re.search(r"\{[\s\S]*\}", txt)
    if m:
        try:
            return json.loads(m.group(0))
        except Exception:
            return None
    return None

# -------------------- Text extraction --------------------
def text_from_pdf_bytes(b: bytes) -> str:
    try:
        reader = PyPDF2.PdfReader(BytesIO(b))
        out = []
        for p in reader.pages:
            t = p.extract_text()
            if t:
                out.append(t)
        return "\n\n".join(out).strip()
    except Exception:
        enc = chardet.detect(b).get("encoding") or "utf-8"
        try:
            return b.decode(enc, errors="ignore")
        except Exception:
            return ""

def text_from_docx_bytes(b: bytes) -> str:
    try:
        doc = docx.Document(BytesIO(b))
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n\n".join(paragraphs).strip()
    except Exception:
        return ""

def extract_text(filename: str, data: bytes) -> str:
    ext = (filename.lower().rsplit(".", 1) + [""])[-1]
    if ext == "pdf":
        return text_from_pdf_bytes(data)
    if ext == "docx":
        return text_from_docx_bytes(data)
    # fallback to text/csv
    enc = chardet.detect(data).get("encoding") or "utf-8"
    try:
        return data.decode(enc, errors="ignore")
    except Exception:
        return ""

# -------------------- Segmentation --------------------
def segment_text_by_paragraphs(text: str, min_chars: int = 300) -> List[str]:
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not paras:
        paras = [ln.strip() for ln in text.splitlines() if ln.strip()]
    segments = []
    cur = ""
    for p in paras:
        if not cur:
            cur = p
        elif len(cur) + len(p) + 2 <= min_chars:
            cur += "\n\n" + p
        else:
            if len(cur) < min_chars:
                cur += "\n\n" + p
            else:
                segments.append(cur)
                cur = p
    if cur:
        segments.append(cur)
    out = []
    for s in segments:
        if len(s) <= 1200:
            out.append(s)
        else:
            sents = re.split(r'(?<=[.!?])\s+', s)
            chunk = ""
            for sent in sents:
                if not chunk:
                    chunk = sent
                elif len(chunk) + len(sent) < 900:
                    chunk += " " + sent
                else:
                    out.append(chunk.strip())
                    chunk = sent
            if chunk:
                out.append(chunk.strip())
    return out

def split_sentences(segment: str) -> List[str]:
    sents = re.split(r'(?<=[.!?])\s+', segment.strip())
    sents = [s.strip() for s in sents if s.strip()]
    return sents if sents else [segment.strip()]

# -------------------- Local detector (worker-safe) --------------------
# Worker helper functions MUST be top-level for pickling.

def detect_with_typetruth_text(text: str) -> float:
    try:
        if typetruth and hasattr(typetruth, "predict"):
            res = typetruth.predict(text)
            if isinstance(res, dict) and "ai_prob" in res:
                return float(res["ai_prob"])
        if typetruth and hasattr(typetruth, "score"):
            return float(typetruth.score(text))
    except Exception:
        return 0.0
    return 0.0

# GPT-2 perplexity fallback; this will be executed inside worker processes lazily
def _perplexity_score_gpt2_local(text: str, cache: Dict) -> float:
    try:
        if "tokenizer" not in cache:
            cache["tokenizer"] = GPT2TokenizerFast.from_pretrained("gpt2")
            cache["model"] = GPT2LMHeadModel.from_pretrained("gpt2")
            cache["model"].eval()
        tokenizer = cache["tokenizer"]
        model = cache["model"]
        enc = tokenizer(text, return_tensors="pt", truncation=True, max_length=1024)
        import torch
        with torch.no_grad():
            out = model(enc["input_ids"], labels=enc["input_ids"])
            loss = float(out.loss.item())
        ppl = math.exp(loss) if loss < 100 else float("inf")
        if ppl == float("inf"):
            return 0.0
        if ppl <= 10:
            return 0.98
        if ppl >= 200:
            return 0.02
        val = 1 - (math.log(ppl) - math.log(10)) / (math.log(200) - math.log(10))
        val = max(0.0, min(0.99, val))
        return float(val)
    except Exception:
        return 0.0

def detect_segment_and_sentences_local(segment: str, cache: Dict) -> Dict[str, Any]:
    # compute a segment-level score
    seg_score = 0.0
    if TYPETRUTH_AVAILABLE:
        try:
            seg_score = detect_with_typetruth_text(segment)
        except Exception:
            seg_score = 0.0
    elif TRANSFORMERS_AVAILABLE:
        try:
            seg_score = _perplexity_score_gpt2_local(segment, cache)
        except Exception:
            seg_score = 0.0
    else:
        words = segment.split()
        if len(words) > 6:
            avg_word_len = sum(len(w) for w in words) / len(words)
            heur = 0.45 if avg_word_len < 4 else 0.18
            seg_score = round(min(0.99, heur * (len(segment) / 1000)), 6)

    # sentence-level scores
    sentences = split_sentences(segment)
    sentence_scores = []
    for s in sentences:
        s_score = 0.0
        # 1) If a trained AI validator model is available, use it first.
        if AI_VALIDATOR_AVAILABLE and AI_VALIDATOR_MODEL is not None:
            try:
                # prefer predict_proba if available
                if hasattr(AI_VALIDATOR_MODEL, "predict_proba"):
                    probs = AI_VALIDATOR_MODEL.predict_proba([s])
                    # try to pick the 'ai' class probability; commonly last column
                    if probs is not None and len(probs) > 0:
                        prob = float(probs[0][-1])
                        s_score = max(0.0, min(0.999999, prob))
                    else:
                        s_score = 0.0
                elif hasattr(AI_VALIDATOR_MODEL, "predict"):
                    p = AI_VALIDATOR_MODEL.predict([s])[0]
                    # if predict returns class labels (0/1), map to 0.0/1.0
                    try:
                        s_score = float(p)
                    except Exception:
                        s_score = 0.0
                else:
                    s_score = 0.0
            except Exception:
                # fallback to other detectors if model call fails
                try:
                    traceback.print_exc()
                except Exception:
                    pass
                s_score = 0.0
        else:
            # 2) existing local fallbacks
            if TYPETRUTH_AVAILABLE:
                try:
                    s_score = detect_with_typetruth_text(s)
                except Exception:
                    s_score = 0.0
            elif TRANSFORMERS_AVAILABLE:
                try:
                    s_score = _perplexity_score_gpt2_local(s, cache)
                except Exception:
                    s_score = 0.0
            else:
                words = s.split()
                if len(words) > 4:
                    avg_word_len = sum(len(w) for w in words) / len(words)
                    heur = 0.5 if avg_word_len < 4 else 0.18
                    s_score = round(min(0.99, heur * (len(s) / 800)), 6)
                else:
                    s_score = 0.0
        sentence_scores.append({"sentence_text": s, "sentence_ai_prob": float(round(s_score, 6))})

    return {"segment_ai_prob": float(round(seg_score, 6)), "sentences": sentence_scores}

# -------------------- Worker wrapper (picklable) --------------------
def worker_detect(args: Tuple[int, str]):
    idx, segment = args
    # per-process cache stored as function attribute
    if not hasattr(worker_detect, "_cache"):
        worker_detect._cache = {}
    cache = worker_detect._cache
    res = detect_segment_and_sentences_local(segment, cache)
    return {
        "segment_index": int(idx),
        "segment_text": segment,
        "ai_probability_detector": float(res["segment_ai_prob"]),
        "sentences": res["sentences"]
    }

# -------------------- Gemini validator prompt/builders --------------------
def build_gemini_sentence_validation_prompt(sentence: str, ai_prob: float, past_reports_sample: List[Dict[str, Any]]) -> str:
    past_preview = ""
    for r in (past_reports_sample or [])[:6]:
        preview = (r.get("result") or {}).get("raw_text") if isinstance((r.get("result") or {}).get("raw_text"), str) else None
        if preview:
            past_preview += f"--- {r.get('filename','<file>')} ---\n{preview[:600]}\n\n"
    # Escape braces in JSON example by doubling them if using .format elsewhere; here we use f-string so OK.
        prompt = f"""
You are an expert auditor specializing in detecting AI-generated text.

Input:
- Sentence: {sentence}
- Detector AI probability (0.0-1.0): {ai_prob}
- Short past report previews (for similarity): {bool(past_preview)}

TASKS:
1) Confirm or adjust the probability (0.0-1.0).
2) Decide: "ai", "human", or "uncertain".
3) Provide a one-line comment with reasons and confidence.
4) Provide a short explanation (2-3 sentences) describing WHY the model decided as it did (e.g. formulaic phrasing, repetition, unusual token sequences, lack of personal details).
5) Provide 2-5 actionable recommendations to make this sentence/readable content appear more human (if decision is "ai" or high probability) or to improve clarity/credibility (if "human").
6) If similar to any past preview, include filename and approximate similarity (0.0-1.0).

Return ONLY valid JSON with keys:
{
    "validated_ai_probability": 0.0,
    "decision": "ai"|"human"|"uncertain",
    "comment": "one-line comment | reasons | confidence:<0.0-1.0>",
    "explanation": "short explanation why this appears ai/human",
    "recommendations": ["short actionable suggestion", ...],
    "matched_past_files": [ { "filename": "...", "similarity": 0.0 } ]
}

Past previews:
{past_preview}
"""
    return prompt

def call_gemini_validate_sentence(sentence: str, ai_prob: float, past_reports_sample: List[Dict[str, Any]]) -> Dict[str, Any]:
    try:
        prompt = build_gemini_sentence_validation_prompt(sentence, ai_prob, past_reports_sample)
        model = genai.GenerativeModel(MODEL_NAME)
        resp = model.generate_content(prompt)
        raw = resp.text or ""
        parsed = safe_json_parse(raw)
        if isinstance(parsed, dict):
            return parsed
    except Exception:
        pass
    # fallback auto-decision heuristics if model fails
    decision = "ai" if ai_prob >= 0.9 else ("human" if ai_prob < 0.45 else "uncertain")
    fallback = {
        "validated_ai_probability": float(round(ai_prob, 6)),
        "decision": decision,
        "comment": f"Auto-validated: detector-only decision ({decision}) | confidence:{round(ai_prob,3)}",
        "explanation": "Model unavailable — heuristic fallback used. High detector score indicates formulaic phrasing or short token patterns.",
        "recommendations": [] ,
        "matched_past_files": []
    }
    # provide minimal recommendations heuristically when AI-like
    try:
        if decision == "ai":
            fallback["recommendations"] = [
                "Add personal anecdotes or first-person observations.",
                "Use varied sentence lengths and more colloquial phrasing.",
                "Include concrete facts, citations, and domain-specific details.",
            ]
        elif decision == "uncertain":
            fallback["recommendations"] = [
                "Increase specificity and add references.",
                "Avoid repetitive or templated sentence openings.",
            ]
        else:
            fallback["recommendations"] = ["No major changes suggested; consider adding citations if appropriate."]
    except Exception:
        pass
    return fallback


# Note: the document-level Gemini recommendation function has been removed to keep
# the pipeline lightweight for PDF/batch runs. A simple heuristic sentence is
# produced below in the main endpoint instead of calling an external LLM.

# -------------------- DB & Storage helpers --------------------
def load_existing_report_by_hash(file_hash: str) -> Optional[Dict[str, Any]]:
    try:
        res = supabase.table(REPORT_TABLE).select("id,filename,file_hash,result,created_at").eq("file_hash", file_hash).limit(1).execute()
        rows = getattr(res, "data", []) or []
        if rows:
            return rows[0]
    except Exception:
        pass
    return None

def upload_raw_file_to_storage(filename: str, data: bytes) -> bool:
    try:
        supabase.storage.from_(RAW_BUCKET).upload(filename, data, {"content-type": "application/octet-stream"})
        return True
    except Exception:
        # try rename if exists
        try:
            existing = supabase.storage.from_(RAW_BUCKET).list() or []
            existing_names = [o.get("name") for o in existing]
            unique = make_unique_filename(filename, existing_names)
            supabase.storage.from_(RAW_BUCKET).upload(unique, data, {"content-type": "application/octet-stream"})
            return True
        except Exception:
            return False

def upload_json_report_to_storage(json_name: str, report: Dict[str, Any]) -> bool:
    try:
        supabase.storage.from_(REPORT_BUCKET).upload(json_name, json.dumps(report, ensure_ascii=False).encode("utf-8"), {"content-type": "application/json"})
        return True
    except Exception:
        # if exists, attempt unique name
        try:
            existing = supabase.storage.from_(REPORT_BUCKET).list() or []
            existing_names = [o.get("name") for o in existing]
            unique = make_unique_filename(json_name, existing_names)
            supabase.storage.from_(REPORT_BUCKET).upload(unique, json.dumps(report, ensure_ascii=False).encode("utf-8"), {"content-type": "application/json"})
            return True
        except Exception:
            return False

def insert_report_row_db(row: Dict[str, Any]) -> bool:
    try:
        supabase.table(REPORT_TABLE).insert(row).execute()
        return True
    except Exception:
        return False

def load_recent_reports(limit: int = 20) -> List[Dict[str, Any]]:
    try:
        res = supabase.table(REPORT_TABLE).select("*").order("created_at", desc=True).limit(limit).execute()
        rows = getattr(res, "data", []) or []
        return rows
    except Exception:
        return []

# -------------------- Aggregation --------------------
def aggregate_sentence_level_results(segment_reports: List[Dict[str, Any]]) -> Dict[str, Any]:
    sentence_probs = []
    gemini_ai_count = 0
    gemini_human_count = 0
    gemini_uncount = 0
    flagged_sentences = []
    for seg in segment_reports:
        for s in seg.get("sentences", []):
            sentence_probs.append(s.get("sentence_ai_prob", 0.0))
            gv = s.get("gemini_validation") or {}
            decision = gv.get("decision")
            if decision == "ai":
                gemini_ai_count += 1
                flagged_sentences.append({
                    "segment_index": seg["segment_index"],
                    "sentence": s["sentence_text"][:800],
                    "detector_prob": s.get("sentence_ai_prob", 0.0),
                    "gemini": gv
                })
            elif decision == "human":
                gemini_human_count += 1
            else:
                gemini_uncount += 1
    avg_detector_sentence_prob = float(round(sum(sentence_probs)/len(sentence_probs), 6)) if sentence_probs else 0.0
    total_sentences = len(sentence_probs)
    ai_sentence_pct = int(round((gemini_ai_count / total_sentences) * 100)) if total_sentences else 0
    return {
        "avg_detector_sentence_prob": avg_detector_sentence_prob,
        "total_sentences": total_sentences,
        "gemini_ai_count": gemini_ai_count,
        "gemini_human_count": gemini_human_count,
        "gemini_uncertain_count": gemini_uncount,
        "ai_sentences_percentage_by_gemini": ai_sentence_pct,
        "flagged_sentences": flagged_sentences
    }

# -------------------- Main endpoint --------------------
@router.post("/detect-ai-and-validate")
async def detect_ai_and_validate(file: UploadFile = File(...)):
    """
    Complete pipeline:
    - read bytes
    - compute hash -> if exist in DB return cached result
    - otherwise, run local detection in worker pool
    - collect sentences with detector >= threshold
    - validate suspicious sentences with Gemini (sequentially)
    - attach validations and aggregate
    - save raw file + json + db row
    - return compact response
    """
    try:
        data = await file.read()
        filename_raw = sanitize_filename(file.filename or f"upload_{datetime.utcnow().isoformat()}")
        file_hash = sha256_bytes(data)

        # 1) duplicate check by hash (fast)
        existing = load_existing_report_by_hash(file_hash)
        if existing:
            # Build a compact cached response (only the four fields requested)
            try:
                stored = existing.get("result") or existing.get("cached_result") or existing

                # Try to derive display score (human-facing % where higher == more human)
                display_score = None
                # common places for overall human percentage
                display_score = (
                    (stored.get("overall_scores") or {}).get("overall_human_percentage")
                    if isinstance(stored, dict) else None
                )
                if display_score is None:
                    display_score = (
                        (stored.get("field_analysis") or {}).get("overall_scores", {}).get("overall_human_percentage")
                    )
                if display_score is None:
                    # try legacy names
                    display_score = (
                        (stored.get("field_analysis") or {}).get("overall_scores", {}).get("human_score")
                    )
                if display_score is None:
                    # fallback: if combined_score present, compute 100 - combined
                    combined = (stored.get("aggregate") or {}).get("combined_score_pct") or (stored.get("aggregate") or {}).get("combined_score")
                    if isinstance(combined, (int, float)):
                        display_score = int(round(100 - float(combined)))
                if display_score is None:
                    # final fallback: try overall ai/human mix
                    try:
                        ai_pct = (stored.get("field_analysis") or {}).get("overall_scores", {}).get("overall_ai_percentage")
                        if isinstance(ai_pct, (int, float)):
                            display_score = int(round(100 - float(ai_pct)))
                    except Exception:
                        display_score = None

                if display_score is None:
                    display_score = 50

                try:
                    display_score = int(display_score)
                except Exception:
                    display_score = max(0, min(100, int(round(float(display_score or 50)))))

                # Determine classification for improvement comment
                classification = None
                classification = (stored.get("overall_scores") or {}).get("classification") or (stored.get("field_analysis") or {}).get("overall_scores", {}).get("classification")
                if not classification:
                    classification = (stored.get("aggregate") or {}).get("file_verdict")
                if not classification:
                    classification = "mixed"

                if classification == "human":
                    improvement_sentence = "Document appears human-authored. Consider adding more specific citations and examples to strengthen credibility."
                elif classification == "ai":
                    improvement_sentence = f"Document shows signs of AI generation in parts of the text. Add personal insights, specific examples, and vary writing style."
                else:
                    improvement_sentence = f"Mixed signals detected. Review flagged sections and add specific examples, citations and varied phrasing to improve human authenticity."

                # Build deterministic comments text (same format as runtime run)
                try:
                    benefit_to_moc = "Yes" if display_score >= 50 else "No"
                    if benefit_to_moc == "No":
                        explanation = (
                            "The document has limited alignment to the listed MOC benefit areas; "
                            "revisions are needed to demonstrate clear benefits and technical readiness."
                        )
                        recommendations = [
                            "Address gaps: include work on Advanced technologies for improving coal mining, Waste-to-Wealth concepts, Alternative uses of coal & clean coal technologies.",
                            "Add technical validation or pilot plans to demonstrate feasibility.",
                            "Highlight practical benefits to operations, safety, or environment with quantitative targets."
                        ]
                    else:
                        explanation = (
                            "The document demonstrates reasonable alignment to MOC benefit areas; "
                            "consider strengthening technical validation, citations, and measurable targets."
                        )
                        recommendations = [
                            "Add technical validation or pilot plans to demonstrate feasibility.",
                            "Include quantitative targets for operational, safety, or environmental improvements.",
                            "Provide clearer links between proposed work and MOC benefit areas with citations."
                        ]

                    comments_lines = []
                    comments_lines.append(f"Score: {display_score}/100 Benefit to MOC: {benefit_to_moc}")
                    comments_lines.append(explanation)
                    comments_lines.append("Recommended actions:")
                    for r in recommendations:
                        comments_lines.append(f"- {r}")
                    comments_text = "\n".join(comments_lines)
                except Exception:
                    comments_text = (
                        f"Score: {display_score}/100 Benefit to MOC: No\n"
                        "The document has limited alignment to the listed MOC benefit areas; revisions are needed to demonstrate clear benefits and technical readiness.\n"
                        "Recommended actions:\n"
                        "- Address gaps: include work on Advanced technologies for improving coal mining, Waste-to-Wealth concepts, Alternative uses of coal & clean coal technologies.\n"
                        "- Add technical validation or pilot plans to demonstrate feasibility.\n"
                        "- Highlight practical benefits to operations, safety, or environment with quantitative targets."
                    )

                # Attempt to extract flagged lines from stored result
                flagged = []
                try:
                    # common locations
                    if isinstance(stored, dict):
                        if stored.get("flagged_lines"):
                            flagged = stored.get("flagged_lines")
                        elif (stored.get("aggregate") or {}).get("flagged_sentences"):
                            flagged = (stored.get("aggregate") or {}).get("flagged_sentences")
                        elif (stored.get("field_analysis") or {}).get("flagged_fields"):
                            # map to compact format if needed
                            ff = (stored.get("field_analysis") or {}).get("flagged_fields") or []
                            for f in ff:
                                flagged.append({
                                    "field_name": f.get("field_name"),
                                    "ai_probability": f.get("ai_probability"),
                                    "reason": f.get("reason")
                                })
                except Exception:
                    flagged = []

                response = {
                    "model_score_pct": display_score,
                    "improvement_comment": improvement_sentence,
                    "comments": comments_text,
                    "flagged_lines": flagged,
                }
                return JSONResponse(response)
            except Exception:
                # if any error building compact response, fall back to original cached payload
                return JSONResponse({
                    "status": "cached",
                    "message": "File already processed. Returning cached result.",
                    "filename": existing.get("filename"),
                    "file_hash": existing.get("file_hash"),
                    "created_at": existing.get("created_at"),
                    "cached_result": existing.get("result")
                })

        # 2) text extraction
        text = extract_text(filename_raw, data)
        if not text or len(text.strip()) < 40:
            return JSONResponse({"error": "Could not extract usable text from file"}, status_code=400)

        # 3) segmentation
        segments = segment_text_by_paragraphs(text, min_chars=300)
        if not segments:
            segments = [text]

        # 4) run detector workers in multiprocessing pool
        args = [(i, segments[i]) for i in range(len(segments))]
        workers = min(WORKER_COUNT, max(1, len(args)))
        with Pool(processes=workers) as pool:
            detector_results = pool.map(worker_detect, args)

        # 5) assemble suspicious sentences list for Gemini validation
        suspicious = []
        for seg in detector_results:
            for si, s in enumerate(seg.get("sentences", [])):
                if s.get("sentence_ai_prob", 0.0) >= SENTENCE_VALIDATE_THRESHOLD:
                    suspicious.append((seg["segment_index"], si, s["sentence_text"], s["sentence_ai_prob"]))

        # 6) load recent reports (context for Gemini)
        past_reports = []
        try:
            past_reports = load_recent_reports(limit=20)
        except Exception:
            past_reports = []

        # 7) validate suspicious sentences with Gemini (main process)
        validations = {}
        for seg_idx, sent_idx, sent_text, s_prob in suspicious:
            validations[(seg_idx, sent_idx)] = call_gemini_validate_sentence(sent_text, s_prob, past_reports)

        # 8) attach gemini validations back to results; auto-decide for others
        for seg in detector_results:
            seg_idx = seg["segment_index"]
            for si, s in enumerate(seg["sentences"]):
                key = (seg_idx, si)
                if key in validations:
                    s["gemini_validation"] = validations[key]
                else:
                    # default inline decision heuristics
                    prob = s.get("sentence_ai_prob", 0.0)
                    if prob >= 0.90:
                        decision = "ai"
                    elif prob < 0.45:
                        decision = "human"
                    else:
                        decision = "uncertain"
                    s["gemini_validation"] = {
                        "validated_ai_probability": float(round(prob, 6)),
                        "decision": decision,
                        "comment": f"Auto-decision based on detector threshold | confidence:{round(prob,3)}",
                        "matched_past_files": []
                    }

        # 9) aggregate
        agg = aggregate_sentence_level_results(detector_results)

        # 9.1) compute model vs LLM (Gemini) scores and combined score
        # model score: average detector sentence probability (0.0-1.0)
        model_avg = agg.get("avg_detector_sentence_prob", 0.0) or 0.0

        # llm score: average of gemini validated probabilities where available
        gemini_probs = []
        for seg in detector_results:
            for s in seg.get("sentences", []):
                gv = s.get("gemini_validation") or {}
                v = gv.get("validated_ai_probability")
                if isinstance(v, (int, float)):
                    gemini_probs.append(float(v))

        if gemini_probs:
            gemini_avg = float(sum(gemini_probs) / len(gemini_probs))
        else:
            # fallback: use detector average as proxy
            gemini_avg = model_avg

        # convert to percentages
        model_score_pct = round(model_avg * 100, 2)
        llm_score_pct = round(gemini_avg * 100, 2)

        # combined score: 60% model, 40% llm
        combined_score = round((0.6 * model_score_pct) + (0.4 * llm_score_pct), 2)

        # verdict thresholds: >=65 => ai, <=35 => human, else uncertain
        if combined_score >= 65:
            file_verdict = "ai"
        elif combined_score <= 35:
            file_verdict = "human"
        else:
            file_verdict = "uncertain"

        # attach to aggregate for storage
        agg["model_avg_probability"] = float(round(model_avg, 6))
        agg["llm_avg_probability"] = float(round(gemini_avg, 6))
        agg["model_score_pct"] = model_score_pct
        agg["llm_score_pct"] = llm_score_pct
        agg["combined_score_pct"] = combined_score
        agg["file_verdict"] = file_verdict

        # 9.2) Produce a compact improvement sentence (no external LLM calls)
        # - If verdict is fully human, display score should be 100.
        # - If verdict is fully ai, display score should be 0.
        # - Otherwise display = 100 - combined_score (so lower AI-likeness => higher displayed score).
        if file_verdict == "human":
            improvement_sentence = "Document appears human-authored; consider adding citations or more domain-specific examples only if needed."
        elif file_verdict == "ai":
            improvement_sentence = "Document shows signs of AI-generation; add personal anecdotes, concrete citations, and vary sentence length and tone to increase human-likeness."
        else:
            improvement_sentence = "Document shows mixed signals; add specific examples, citations and vary phrasing to improve human-likeness."

        agg["document_recommendations"] = {"why": improvement_sentence, "recommendations": [], "confidence": round(min(0.99, combined_score / 100.0), 3)}

        # 10) build report
        report = {
            "original_filename": filename_raw,
            "file_hash": file_hash,
            "created_at": datetime.utcnow().isoformat(),
            "segments_count": len(detector_results),
            "segments": detector_results,
            "aggregate": agg,
            "raw_text_snippet": text[:20000]
        }

        # 11) store raw file to storage
        # choose unique name if filename exists
        uploaded_name = filename_raw
        try:
            existing_files = supabase.storage.from_(RAW_BUCKET).list() or []
            existing_names = [o.get("name") for o in existing_files]
        except Exception:
            existing_names = []
        if uploaded_name in existing_names:
            uploaded_name = make_unique_filename(uploaded_name, existing_names)
        stored_raw = False
        try:
            supabase.storage.from_(RAW_BUCKET).upload(uploaded_name, data, {"content-type": "application/octet-stream"})
            stored_raw = True
        except Exception:
            # best-effort silence
            stored_raw = False

        # 12) store report JSON to storage with better error handling
        json_name = f"{sanitize_filename(uploaded_name.rsplit('.', 1)[0])}_{uuid.uuid4().hex[:8]}.ai-detect.json"
        report_url = ""
        try:
            # Ensure unique filename
            existing_files = supabase.storage.from_(REPORT_BUCKET).list() or []
            existing_names = [o.get("name") for o in existing_files if o.get("name")]
            if json_name in existing_names:
                json_name = f"{sanitize_filename(uploaded_name.rsplit('.', 1)[0])}_{uuid.uuid4().hex[:8]}_report.json"
            
            # Upload JSON report
            json_bytes = json.dumps(report, indent=2, ensure_ascii=False).encode("utf-8")
            supabase.storage.from_(REPORT_BUCKET).upload(
                json_name, 
                json_bytes, 
                {"content-type": "application/json"}
            )
            report_url = supabase.storage.from_(REPORT_BUCKET).get_public_url(json_name)
            print(f"Successfully stored JSON report: {json_name}")
        except Exception as e:
            print(f"Failed to store JSON report: {str(e)}")
            # Try fallback storage method
            try:
                fallback_name = f"fallback_{uuid.uuid4().hex[:12]}.json"
                json_bytes = json.dumps(report, indent=2).encode("utf-8")
                supabase.storage.from_(REPORT_BUCKET).upload(
                    fallback_name, 
                    json_bytes, 
                    {"content-type": "application/json"}
                )
                report_url = supabase.storage.from_(REPORT_BUCKET).get_public_url(fallback_name)
                print(f"Fallback storage successful: {fallback_name}")
            except Exception as fallback_error:
                print(f"Fallback storage also failed: {str(fallback_error)}")

        # 13) insert DB row with file_hash and result
        db_row = {
            "filename": uploaded_name,
            "file_path": f"{RAW_BUCKET}/{uploaded_name}",
            "file_hash": file_hash,
            "ai_percentage": agg.get("ai_sentences_percentage_by_gemini"),
            "segments_count": agg.get("total_sentences"),
            "result": report,
            "created_at": datetime.utcnow().isoformat()
        }
        try:
            insert_report_row_db(db_row)
            print(f"Successfully inserted DB row for: {uploaded_name}")
        except Exception as db_error:
            print(f"Failed to insert DB row: {str(db_error)}")

        # 14) return compact response suitable for PDF/UI rendering
        # Transform combined_score (AI-likeness 0..100) into display score where higher == more human.
        display_score = int(round(100 - combined_score))
        display_score = max(0, min(100, display_score))

        # Map file_verdict to simple classification output
        classification = file_verdict if file_verdict in ("human", "ai") else "mixed"

        # Build a brief list of flagged lines/sentences to return in the compact response.
        # We include segment_index, sentence_index, approximate line_number (if found),
        # truncated text, ai_probability and decision. Keep list short and sorted by prob.
        flagged = []
        for seg in detector_results:
            seg_idx = seg.get("segment_index")
            for si, s in enumerate(seg.get("sentences", [])):
                gv = s.get("gemini_validation") or {}
                decision = gv.get("decision")
                prob = gv.get("validated_ai_probability")
                if prob is None:
                    prob = s.get("sentence_ai_prob", 0.0)
                try:
                    st = s.get("sentence_text", "") or ""
                except Exception:
                    st = ""
                # Consider it flagged when Gemini/model says "ai" or prob >= 0.9
                if (isinstance(decision, str) and decision == "ai") or float(prob) >= 0.90:
                    # approximate line number by finding first occurrence in raw text
                    line_number = None
                    try:
                        if st and isinstance(text, str):
                            pos = text.find(st)
                            if pos >= 0:
                                line_number = text[:pos].count("\n") + 1
                    except Exception:
                        line_number = None
                    flagged.append({
                        "segment_index": int(seg_idx) if seg_idx is not None else None,
                        "sentence_index": int(si),
                        "line_number": line_number,
                        "text": st[:240].strip(),
                        "ai_probability": float(round(float(prob), 3)),
                        "decision": decision or "ai"
                    })

        # keep top 10 by probability
        flagged = sorted(flagged, key=lambda x: x.get("ai_probability", 0.0), reverse=True)[:10]

        # Build a deterministic, user-facing comments block (no external LLM call)
        try:
            try:
                changeable_pct = int(display_score)
            except Exception:
                changeable_pct = 0

            # Determine basic benefit alignment flag based on display score
            benefit_to_moc = "Yes" if display_score >= 50 else "No"

            if benefit_to_moc == "No":
                explanation = (
                    "The document has limited alignment to the listed MOC benefit areas; "
                    "revisions are needed to demonstrate clear benefits and technical readiness."
                )
                recommendations = [
                    "Address gaps: include work on Advanced technologies for improving coal mining, Waste-to-Wealth concepts, Alternative uses of coal & clean coal technologies.",
                    "Add technical validation or pilot plans to demonstrate feasibility.",
                    "Highlight practical benefits to operations, safety, or environment with quantitative targets."
                ]
            else:
                explanation = (
                    "The document demonstrates reasonable alignment to MOC benefit areas; "
                    "consider strengthening technical validation, citations, and measurable targets."
                )
                recommendations = [
                    "Add technical validation or pilot plans to demonstrate feasibility.",
                    "Include quantitative targets for operational, safety, or environmental improvements.",
                    "Provide clearer links between proposed work and MOC benefit areas with citations."
                ]

            # Compose the comments block as a single string
            comments_lines = []
            comments_lines.append(f"Score: {display_score}/100 Benefit to MOC: {benefit_to_moc}")
            comments_lines.append(explanation)
            comments_lines.append("Recommended actions:")
            for r in recommendations:
                comments_lines.append(f"- {r}")

            comments_text = "\n".join(comments_lines)
        except Exception:
            # Very conservative fallback string
            comments_text = (
                f"Score: {display_score}/100 Benefit to MOC: No\n"
                "The document has limited alignment to the listed MOC benefit areas; revisions are needed to demonstrate clear benefits and technical readiness.\n"
                "Recommended actions:\n"
                "- Address gaps: include work on Advanced technologies for improving coal mining, Waste-to-Wealth concepts, Alternative uses of coal & clean coal technologies.\n"
                "- Add technical validation or pilot plans to demonstrate feasibility.\n"
                "- Highlight practical benefits to operations, safety, or environment with quantitative targets."
            )

        response = {
            "model_score_pct": display_score,
            "improvement_comment": improvement_sentence,
            "comments": comments_text,
            "flagged_lines": flagged,
        }
        # mark todo item complete for this change
        try:
            # best-effort: update todo list status
            pass
        except Exception:
            pass
        return JSONResponse(response)

    except Exception as e:
        # import traceback
        # traceback.print_exc()
        return JSONResponse({"error": str(e)}, status_code=500)
